#web_search.py
import json
import re
import sys
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from pathlib import Path

import requests

def _get_base_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).resolve().parent.parent


BASE_DIR        = _get_base_dir()
API_CONFIG_PATH = BASE_DIR / "config" / "api_keys.json"


def _get_api_key() -> str:
    with open(API_CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)["gemini_api_key"]


def _gemini_search(query: str) -> str:
    from google import genai

    client   = genai.Client(api_key=_get_api_key())
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=query,
        config={"tools": [{"google_search": {}}]},
    )

    text = ""
    for part in response.candidates[0].content.parts:
        if hasattr(part, "text") and part.text:
            text += part.text

    text = text.strip()
    if not text:
        raise ValueError("Gemini returned an empty response.")
    return text


def _ddg_search(query: str, max_results: int = 6) -> list[dict]:
    try:
        from ddgs import DDGS
    except ImportError:
        from duckduckgo_search import DDGS

    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=max_results):
            results.append({
                "title":   r.get("title",  ""),
                "snippet": r.get("body",   ""),
                "url":     r.get("href",   ""),
            })
    return results


_PH_NEWS_FEEDS = (
    ("Rappler",   "https://www.rappler.com/feed/"),
    ("Inquirer",  "https://newsinfo.inquirer.net/feed"),
)

_PH_NEWS_HINTS = (
    "news", "headline", "headlines", "rappler", "inquirer",
    "philippine", "ph news", "latest story", "what happened",
    "briefing", "top story",
)


def _is_ph_news_query(query: str) -> bool:
    q = query.lower()
    return any(h in q for h in _PH_NEWS_HINTS)


def _parse_rss_headlines(source: str, url: str, limit: int = 4) -> list[dict]:
    items: list[dict] = []
    try:
        resp = requests.get(
            url,
            timeout=12,
            headers={"User-Agent": "Jarvis-MarkXXXIX/1.0"},
        )
        resp.raise_for_status()
        root = ET.fromstring(resp.content)
        for node in root.findall(".//item")[:limit]:
            title = (node.findtext("title") or "").strip()
            link  = (node.findtext("link") or "").strip()
            pub   = (node.findtext("pubDate") or node.findtext("published") or "").strip()
            if title:
                items.append({"source": source, "title": title, "link": link, "pub": pub})
    except Exception as e:
        print(f"[WebSearch] ⚠️ RSS {source} failed: {e}")
    return items


def _fetch_ph_headlines() -> str:
    """Latest headlines from Rappler and Inquirer RSS feeds."""
    all_items: list[dict] = []
    for source, url in _PH_NEWS_FEEDS:
        all_items.extend(_parse_rss_headlines(source, url, limit=4))

    if not all_items:
        try:
            return _gemini_search(
                "List the latest top headlines today from rappler.com and inquirer.net only. "
                "Return exactly 6 items: 3 from Rappler, 3 from Inquirer. "
                "Format each as: [SOURCE] Headline — one-line summary."
            )
        except Exception as e:
            print(f"[WebSearch] ⚠️ PH headline fallback failed: {e}")
            results = _ddg_search(
                "site:rappler.com OR site:inquirer.net Philippines news today headlines",
                max_results=8,
            )
            return _format_ddg("Philippines headlines — Rappler & Inquirer", results)

    lines = [
        "PHILIPPINE HEADLINES — Rappler & Inquirer",
        f"Fetched {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        "",
    ]
    for i, item in enumerate(all_items[:8], 1):
        src = item.get("source", "News")
        title = re.sub(r"\s+", " ", item.get("title", "")).strip()
        pub = item.get("pub", "")
        if pub:
            lines.append(f"{i}. [{src}] {title}")
            lines.append(f"   · {pub}")
        else:
            lines.append(f"{i}. [{src}] {title}")
        link = item.get("link", "")
        if link:
            lines.append(f"   {link}")
        lines.append("")
    return "\n".join(lines).strip()


def _format_ddg(query: str, results: list[dict]) -> str:
    if not results:
        return f"No results found for: {query}"

    lines = [f"Search results for: {query}\n"]
    for i, r in enumerate(results, 1):
        if r.get("title"):   lines.append(f"{i}. {r['title']}")
        if r.get("snippet"): lines.append(f"   {r['snippet']}")
        if r.get("url"):     lines.append(f"   {r['url']}")
        lines.append("")
    return "\n".join(lines).strip()

def _compare(items: list[str], aspect: str) -> str:
    query = (
        f"Compare {', '.join(items)} in terms of {aspect}. "
        "Give specific facts and data."
    )
    try:
        return _gemini_search(query)
    except Exception as e:
        print(f"[WebSearch] ⚠️ Gemini compare failed: {e} — falling back to DDG")

    # DDG fallback: fetch results per item and merge
    all_results: dict[str, list] = {}
    for item in items:
        try:
            all_results[item] = _ddg_search(f"{item} {aspect}", max_results=3)
        except Exception:
            all_results[item] = []

    lines = [f"Comparison — {aspect.upper()}", "─" * 40]
    for item in items:
        lines.append(f"\n▸ {item}")
        for r in all_results.get(item, [])[:2]:
            if r.get("snippet"):
                lines.append(f"  • {r['snippet']}")
    return "\n".join(lines)

def web_search(
    parameters:     dict,
    response=None,
    player=None,
    session_memory=None,
) -> str:
    params = parameters or {}
    query  = params.get("query", "").strip()
    mode   = params.get("mode",  "search").lower().strip()
    if mode in ("ph_headlines", "ph_news", "headlines"):
        mode = "ph_headlines"
    items  = params.get("items", [])
    aspect = params.get("aspect", "general").strip() or "general"
    save_as_docx = params.get("save_as_docx", False)
    filename = params.get("filename", "search_results.docx")

    if not query and not items:
        return "Please provide a search query, sir."

    if items and mode != "compare":
        mode = "compare"

    if player:
        player.write_log(f"[Search] {query or ', '.join(items)}")

    print(f"[WebSearch] 🔍 Query: {query!r}  Mode: {mode}")

    try:
        if mode == "ph_headlines" or _is_ph_news_query(query):
            print("[WebSearch] 📰 Fetching Rappler + Inquirer headlines...")
            result = _fetch_ph_headlines()
            print("[WebSearch] ✅ PH headlines ready.")
        elif mode == "compare" and items:
            print(f"[WebSearch] 📊 Comparing: {items}")
            result = _compare(items, aspect)
            print("[WebSearch] ✅ Compare done.")
        else:
            print("[WebSearch] 🌐 Trying Gemini...")
            try:
                result = _gemini_search(query)
                print("[WebSearch] ✅ Gemini OK.")
            except Exception as e:
                print(f"[WebSearch] ⚠️ Gemini failed ({e}) — trying DDG...")
                results = _ddg_search(query)
                result  = _format_ddg(query, results)
                print(f"[WebSearch] ✅ DDG: {len(results)} result(s).")

        # Save to docx if requested
        if save_as_docx and result:
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                from pathlib import Path
                
                output_dir = Path(r"C:\Users\jaspe\OneDrive\Documents\Jarvis")
                output_dir.mkdir(parents=True, exist_ok=True)
                output_path = output_dir / filename
                if not filename.endswith(".docx"):
                    output_path = output_path.with_suffix(".docx")
                
                doc = Document()
                
                # Add main title with bold formatting
                title = doc.add_heading(filename.replace(".docx", ""), 0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in title.runs:
                    run.font.size = Pt(18)
                    run.font.bold = True
                
                # Process content with better formatting
                lines = result.strip().split('\n')
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Detect headings (lines that end with colon or are short and uppercase)
                    if line.endswith(':') or (len(line) < 50 and line.isupper()):
                        # Add as heading with bold
                        heading = doc.add_heading(line, level=1)
                        for run in heading.runs:
                            run.font.size = Pt(16)
                            run.font.bold = True
                    # Detect bullet points or numbered lists
                    elif line.startswith(('•', '-', '*', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                        paragraph = doc.add_paragraph(line, style='List Bullet')
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                            run.font.bold = True
                    # Regular paragraph - bold key terms before colon
                    else:
                        if ':' in line:
                            parts = line.split(':', 1)
                            if len(parts) == 2:
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run(parts[0] + ':')
                                run.font.size = Pt(12)
                                run.font.bold = True
                                run = paragraph.add_run(parts[1])
                                run.font.size = Pt(12)
                            else:
                                paragraph = doc.add_paragraph(line)
                                for run in paragraph.runs:
                                    run.font.size = Pt(12)
                        else:
                            paragraph = doc.add_paragraph(line)
                            for run in paragraph.runs:
                                run.font.size = Pt(12)
                
                doc.save(output_path)
                print(f"[WebSearch] 📄 Saved to: {output_path}")
                return f"{result}\n\nResearch saved to: {output_path}"
            except Exception as e:
                print(f"[WebSearch] ⚠️ Failed to save docx: {e}")
                return result
        
        return result

    except Exception as e:
        print(f"[WebSearch] ❌ All backends failed: {e}")
        return f"Search failed, sir: {e}"